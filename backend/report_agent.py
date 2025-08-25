from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from typing import List, Union, Dict, Any
from string import Template
import cohere
from docx import Document
from fpdf import FPDF
import io
import re
import os
from fastapi.responses import StreamingResponse
from datetime import datetime, timedelta

# Initialize FastAPI and Cohere
app = FastAPI()
COHERE_API_KEY = os.getenv("COHERE_API_KEY")
cohere_client = cohere.Client(COHERE_API_KEY)

# Pydantic Models
class TextBlock(BaseModel):
    text: str

class DocumentRequest(BaseModel):
    form_data: Dict[str, Any]

class ReportOutput(BaseModel):
    title: str
    blocks: List[Union[TextBlock]]
    metadata: Dict[str, Any]

# Custom PDF Class
class LeaseAgreementPDF(FPDF):
    def __init__(self, title="Lease Agreement"):
        super().__init__()
        self.title = title

    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, self.title, align='C', ln=True)
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')

# Comprehensive lease template
LEASE_TEMPLATE = """
STATE OF {state}                                                    Rev. {revision_date}

LEASE AGREEMENT

This Lease Agreement (this "Agreement") is made this {current_date}, by and between {landlord_name} ("Landlord") and {tenant_names} ("Tenant"). Each Landlord and Tenant may be referred to individually as a "Party" and collectively as the "Parties."

1. PREMISES
The premises leased is a {property_type} located at {property_address}, {property_city}, {property_state} {property_zipcode} (the "Premises").

2. TERM
This Agreement will be for a term beginning on {start_date} and ending on {end_date} (the "Term").

3. RENT AND PAYMENTS
3.1 Monthly Rent: ${monthly_rent}
    Payable in advance on the 1st day of each month.
3.2 Security Deposit: ${security_deposit}
    To be returned within 30 days of lease termination, less any deductions.
3.3 Late Fees: 5% of monthly rent if paid after the 5th day of the month.

4. UTILITIES AND SERVICES
Tenant is responsible for the following utilities:
- Electricity
- Water/Sewer
- Gas
- Internet/Cable
- Trash Collection

5. OCCUPANCY AND USE
5.1 The Premises shall be occupied only by the named Tenant(s).
5.2 The Premises shall be used as a residential dwelling only.

6. MAINTENANCE AND REPAIRS
6.1 Tenant Responsibilities:
    - Keep the Premises clean and sanitary
    - Proper use of all appliances and fixtures
    - Promptly report any maintenance issues
6.2 Landlord Responsibilities:
    - Maintain structural elements
    - Ensure working plumbing, heating, and electrical systems
    - Provide pest control as needed

7. RULES AND REGULATIONS
7.1 Noise and Disturbances: Quiet hours 10 PM - 8 AM
7.2 Smoking: {smoking_policy}
7.3 Pets: {pet_policy}
7.4 Alterations: No alterations without written consent

8. ADDITIONAL TERMS
{additional_terms}

9. DEFAULT AND REMEDIES
9.1 Events of Default:
    - Failure to pay rent when due
    - Violation of any term of this Agreement
9.2 Remedies:
    - Termination of lease
    - Legal action for unpaid rent
    - Eviction proceedings

10. SIGNATURES AND DATE

_______________________
Landlord Signature
Date: {current_date}

_______________________
Tenant Signature
Date: {current_date}

"""

@app.post("/generate_report", response_model=ReportOutput)
async def generate_report(request: DocumentRequest):
    try:
        form_data = request.form_data
        current_date = datetime.now().strftime("%B %d, %Y")
        
        # Prepare the comprehensive prompt for Cohere
        comprehensive_prompt = f"""
Generate a detailed, professional lease agreement using these details:

LANDLORD INFO:
Name: {form_data.get('landlord_name')}
Address: {form_data.get('landlord_address')}
Location: {form_data.get('landlord_city')}, {form_data.get('landlord_state')} {form_data.get('landlord_zipcode')}

PROPERTY INFO:
Type: {form_data.get('property_type')}
Address: {form_data.get('property_address')}
Monthly Rent: ${form_data.get('monthly_rent')}
Security Deposit: ${form_data.get('security_deposit')}

LEASE TERM:
Start: {form_data.get('start_date')}
End: {form_data.get('end_date')}

Create a comprehensive lease agreement that includes all standard protections and clauses.
Use formal legal language but ensure it remains clear and understandable.
Include sections for utilities, maintenance, default conditions, and remedies.
"""

        # Generate using Cohere
        response = cohere_client.generate(
            model="command-xlarge",
            prompt=comprehensive_prompt,
            max_tokens=2500,
            temperature=0.4,
            stop_sequences=["\n\n\n"]
        )
        
        generated_text = response.generations[0].text.strip()
        
        # Fallback to template if generation is insufficient
        if len(generated_text) < 1000:
            # Set default policies
            smoking_policy = "Smoking is strictly prohibited inside the Premises."
            pet_policy = "No pets allowed without prior written consent from Landlord."
            
            # Fill template
            generated_text = LEASE_TEMPLATE.format(
                current_date=current_date,
                revision_date=datetime.now().strftime("%Y%m%d"),
                smoking_policy=smoking_policy,
                pet_policy=pet_policy,
                **form_data
            )

        return ReportOutput(
            title="Lease Agreement",
            blocks=[TextBlock(text=generated_text)],
            metadata={
                "generated_by": "LegalAI",
                "timestamp": datetime.now().isoformat(),
                "version": "2.0"
            }
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate_pdf")
async def generate_pdf(request: DocumentRequest):
    try:
        # Get the generated report first
        report = await generate_report(request)
        report_text = report.blocks[0].text

        # Create PDF
        pdf = LeaseAgreementPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Split text into lines and add to PDF
        for line in report_text.split('\n'):
            if line.strip().upper() == line.strip():  # Check if line is a header
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, line, ln=True)
                pdf.set_font("Arial", size=12)
            else:
                pdf.multi_cell(0, 10, line)

        # Create PDF in memory
        pdf_bytes = io.BytesIO()
        pdf_bytes.write(pdf.output(dest='S').encode('latin-1'))
        pdf_bytes.seek(0)

        return StreamingResponse(
            pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=lease_agreement.pdf"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate_docx")
async def generate_docx(request: DocumentRequest):
    try:
        # Get the generated report first
        report = await generate_report(request)
        report_text = report.blocks[0].text

        # Create DOCX
        doc = Document()
        doc.add_heading('Lease Agreement', 0)

        # Add content to document
        for paragraph in report_text.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        # Create DOCX in memory
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        return StreamingResponse(
            docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=lease_agreement.docx"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))